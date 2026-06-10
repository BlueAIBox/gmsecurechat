from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.document import Document as DocumentClass
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = max(
    [p for p in ROOT.glob("*.docx") if not p.name.startswith("~$") and p.name != "报告_修改版.docx"],
    key=lambda p: p.stat().st_size,
)
OUT_PATH = DOCX_PATH.with_name(f"{DOCX_PATH.stem}_修改版.docx")
BACKUP_PATH = DOCX_PATH.with_name(f"{DOCX_PATH.stem}_修改前备份.docx")
ARCH_IMAGE = ROOT / "系统架构.png"


SECTION_ORDER = [
    ("发送消息", 1),
    ("群组功能", 2),
    ("部署合约", 3),
    ("验证钱包", 4),
    ("转账", 5),
    ("双签名认证", 6),
    ("消息篡改以及文件篡改", 7),
    ("核验存证", 8),
    ("接入agent", 9),
    ("共享agent", 10),
]
SECTION_NUM = dict(SECTION_ORDER)
SECTION_TITLE = {v: k for k, v in SECTION_ORDER}


SECTION_EXPLANATIONS = {
    1: [
        "本节展示的是局域网单聊的完整安全链路。用户登录后会通过 UDP 广播交换昵称、IP、SM2 公钥和可选钱包地址，发送方选中接收方后，并不是直接把明文发送出去，而是先对消息明文计算 SM3 摘要，再用发送方 SM2 私钥对摘要签名，随后随机生成一次性 SM4 会话密钥和 IV，对正文进行 SM4-CBC 加密，最后再用接收方 SM2 公钥封装该会话密钥。",
        "接收方收到 UDP 密文包后，先用自己的 SM2 私钥解封 SM4 会话密钥，再解密得到明文，随后重新计算 SM3 摘要并用发送方 SM2 公钥验签。只有 SM4 解密、SM3 完整性校验和 SM2 身份验签均通过，消息才会进入聊天框。因此聊天界面显示的是用户可读内容，安全日志则保留了密钥封装、摘要、签名、验签等中间过程，便于答辩时逐步说明国密算法如何参与通信。",
    ],
    2: [
        "群组功能不是简单的公共广播明文，而是在原有单聊安全链路基础上做了组消息分发。创建群组时，用户从在线列表中选择多个成员并命名群组，系统保存群组 ID、群名和成员 IP；发送群消息时，程序会对每个群成员分别构造一份国密加密包，逐个使用该成员的 SM2 公钥封装一次性 SM4 会话密钥。",
        "这种设计避免了所有成员共用同一份明文密钥，也能让每个接收端独立完成 SM2 解封、SM4 解密、SM3 完整性校验和 SM2 验签。即使群聊界面看起来像一次群发，底层仍然是多条独立的点对点国密安全消息链路，更适合解释为“基于 P2P 安全通道的轻量群组通信”。",
    ],
    3: [
        "合约部署部分用于把本地审计证据扩展到 Sepolia 测试链。项目中的 GMAuditLedger 智能合约并不保存聊天明文、文件内容或私钥，而是保存经过 SM3/Merkle/会话哈希链处理后的摘要型证据，例如记录类型、SM3 摘要、文件 Merkle Root、当前会话链头和元数据。这样既能体现区块链不可篡改的存证能力，又不会把敏感通信内容暴露到链上。",
        "部署成功后，后续的双签名绑定、文件 Merkle 存证、转账审计和篡改实验记录都可以通过 DApp 提交到合约。答辩时可以强调：链上合约负责“证明某个安全事件曾经发生且摘要未被更改”，国密算法负责“本地通信过程中的加密、签名、摘要和完整性校验”。",
    ],
    4: [
        "钱包验证解决的是“局域网用户身份”和“链上钱包身份”是否属于同一人的问题。由于局域网昵称可能重名，仅凭昵称或 IP 不能证明对方就是某个 MetaMask 地址的控制者，因此系统引入了双重证明：一方面使用对方通信身份的 SM2 私钥对绑定声明签名，另一方面要求对方使用 MetaMask 对同一份绑定声明进行钱包签名。",
        "验证发起方会把对方昵称、对方 SM2 公钥指纹、钱包地址、请求方 SM2 指纹、nonce 和用途组合成绑定声明，并通过国密加密通道发送给对方。对方确认后在 DApp 中调用 MetaMask 签名，exe 再把钱包签名和 SM2 签名通过加密通道回传。发起方同时验证 SM2 签名和钱包签名，只有两者都指向同一身份关系时，列表中才显示钱包已验证。",
    ],
    5: [
        "转账功能把局域网聊天对象和 Sepolia 测试链钱包结合起来。用户在聊天界面选择收款人并输入金额后，exe 会先检查对方钱包地址以及钱包绑定验证状态，再打开本地 DApp 调用 MetaMask 发起 SepoliaETH 测试币转账。真正的链上资产转移由 MetaMask 和以太坊测试网完成，exe 不保存钱包私钥，也不能绕过用户确认。",
        "转账成功后，系统会把交易哈希、发送方、接收方、金额、会话链头等信息写入聊天事件和审计记录，必要时再提交到智能合约。这样聊天框不仅能发送文字和文件，也能展示“安全通信上下文中的链上转账事件”，体现通信系统向可信交易协同平台扩展的能力。",
    ],
    6: [
        "双签名认证用于证明一个身份绑定事件同时经过通信身份和链上身份确认。传统聊天工具通常只证明“我拥有这个聊天账号”，而本项目进一步把 SM2 通信密钥、链上身份材料和绑定声明放在同一个闭环中验证：通信侧使用 SM2 签名保证局域网身份不可伪造，链上侧使用钱包签名或模拟链上 SM2 身份保证钱包控制权可验证。",
        "该功能的价值在于把“人、通信公钥、钱包地址、绑定声明、nonce”组织成可回放校验的证据结构。即使存在同名用户，也必须同时通过 SM2 私钥和钱包签名，才能证明当前局域网用户与指定链上地址具有关联关系。",
    ],
    7: [
        "消息篡改和文件篡改用于演示安全机制如何发现异常。消息篡改现在接入真实发送链路：点击消息篡改后，下一条真实文字消息会先正常完成 SM3 摘要、SM2 签名、SM4 加密和 SM2 密钥封装，然后程序在 UDP 发出前故意改写消息包中的 SM3 字段。接收方收到后会解密并重新计算摘要，发现本地 SM3 与收到的 SM3 不一致，最终拒绝将该消息显示到聊天框。",
        "文件篡改重点展示 SM3 文件哈希与 Merkle Root 的完整性检测。系统会记录原始文件数据、原始 SM3、原始 Merkle Root，再构造被修改后的文件数据并计算篡改后 SM3/Merkle Root。只要文件任意字节发生变化，摘要和 Merkle 根都会不同，日志中会明确显示篡改被发现。该部分适合说明：消息更强调 SM3+SM2 的实时验签拦截，文件更强调 ZUC 传输加密、SM3 摘要和 Merkle 树存证的组合校验。",
    ],
    8: [
        "核验存证功能用于验证本地审计账本是否被修改。系统会把关键安全事件写入 jsonl 账本，每条记录包含 prev、payload、timestamp、recordType 和 recordHash，其中 recordHash 由上一条链头和当前记录内容计算得到。这样每条记录都依赖前一条记录，形成国密会话哈希链。",
        "当用户导入本地存证文件并执行核验时，程序会从第一条记录开始重新计算链头。如果中间任意一个字段被修改，例如截图中把某个数字从 173 改为 170，那么该条记录以及后续所有链头都会失配，验证结果变为 false。该功能证明了项目不仅能发现通信过程中的篡改，也能发现事后对审计证据文件的篡改。",
    ],
    9: [
        "Hermes Agent 是系统中的本机智能安全助手节点。它通过 hermes -z 的 CLI 方式接入，不参与 A/B 之间的密钥交换，不保存双方会话密钥，也不替代 SM2、SM3、SM4、ZUC 的国密安全流程。用户与 Hermes 对话时，系统先在本地完成国密流水线记录，再调用 WSL 中的 Hermes CLI 获得智能回复。",
        "该功能体现了国密安全通信与现代 AI Agent 的融合创新。Hermes 可以辅助解释协议流程、分析 SM 算法日志、总结审计文件、生成答辩说明或帮助定位运行问题，相当于系统内置的“安全通信讲解员”和“智能运维助手”。",
    ],
    10: [
        "共享 Agent 功能进一步把本机智能助手扩展为可授权的局域网服务。A 端本机运行 Hermes，B/C 端不需要安装 WSL 或 Hermes 环境；当 A 选择某个真实在线用户并点击共享 Agent 时，系统会先通过 SM2 身份挑战、确认和加密控制消息完成授权，只有通过身份确认的对方才会在联系人列表中看到 A 共享出来的 Hermes 代理联系人。",
        "共享成功后，对方发送给 Hermes 代理的消息仍会先走国密通信通道到 A 端，再由 A 端本机 Hermes 处理并把回复加密返回。这样既体现 Agent 服务共享能力，又避免把本机环境依赖直接暴露给所有局域网用户。",
    ],
}


ARCH_EXPLANATION = [
    "系统总体采用“五层深度防御体系”：用户层负责 A/B/C 客户端和本机 Hermes Agent 的人机交互；通信层负责 UDP 局域网发现、单聊/群聊分发以及文件传输；国密安全层把 SM2、SM3、SM4、ZUC 和 Merkle Root 分别嵌入身份认证、摘要校验、消息加密、文件流加密和完整性证明；审计与实验层负责算法中间值追踪、国密会话哈希链、消息/文件篡改实验以及本地账本核验；链上扩展层通过 Web DApp、MetaMask、Sepolia 测试链和 GMAuditLedger 合约完成摘要型证据上链。",
    "该架构的核心思想是：通信内容在本地端到端完成国密保护，链上只保存不可逆摘要和审计元数据，不上传聊天明文和文件内容。A/B/C 之间每条消息均执行 SM3 摘要、SM2 签名、SM4-CBC 加密和 SM2 密钥封装；文件传输则使用 ZUC 分片加密并结合 SM3/Merkle Root 做完整性存证；DApp 与智能合约承担可信存证和转账审计角色；Hermes Agent 作为本机智能助手提供协议解释与日志分析能力。",
]


SUMMARY_PARAGRAPHS = [
    "本项目围绕“基于国密算法的安全通信”这一主题，完成了从基础加密通信到审计存证、区块链扩展和 AI Agent 融合的综合设计。系统以局域网 A/B/C 客户端为核心通信场景，重点展示 SM2、SM3、SM4、ZUC 在聊天和文件传输中的真实使用流程：SM2 用于身份签名、验签和会话密钥封装，SM3 用于消息摘要、文件摘要和审计哈希链，SM4-CBC 用于文字消息加密，ZUC 用于文件分片流加密，Merkle Root 用于文件完整性证明。",
    "在功能实现层面，系统已经具备单聊安全通信、群组通信、文件传输、逐步加密流水线日志、安全态势感知面板、双签名身份绑定、对方钱包归属验证、Sepolia 测试币转账、智能合约存证、消息/文件篡改实验台、核验存证以及 Hermes Agent 智能辅助等能力。相比普通加密聊天程序，本项目不只停留在“能发消息”，而是把通信前的身份确认、通信中的加密签名、通信后的日志追踪、异常情况下的篡改检测以及链上审计闭环串联起来。",
    "项目的创新性主要体现在三个方面：第一，将国密算法的中间值和验证结果可视化输出，使用户能够清楚看到每条消息、每个文件如何经过 SM3、SM2、SM4/ZUC 的处理；第二，将本地审计日志、Merkle 树存证和 Sepolia 智能合约结合，使安全事件可以被追溯、核验和上链证明；第三，引入 Hermes Agent 作为本机智能安全助手，在不破坏原有 A/B 国密通信模型的前提下，提供协议解释、日志分析和答辩辅助能力，体现传统密码应用向智能安全通信平台演进的方向。",
    "从安全目标看，系统实现了身份真实性、数据机密性、消息完整性、文件完整性、行为可追溯和存证不可抵赖等多重目标。正常通信时，接收方只有在 SM4/ZUC 解密、SM3 摘要校验和 SM2 验签均通过后才放行内容；发生消息篡改、文件字节篡改或审计文件被修改时，系统能够在日志和核验结果中体现失败原因，并阻止异常内容进入正常聊天视图。",
    "后续可以继续优化的方向包括：进一步完善真实网络攻击模拟，把文件篡改也扩展为“下一次真实文件发送攻击模式”；在智能合约中增加更细粒度的事件索引，便于按用户、会话和记录类型查询；对群组通信引入更正式的群密钥协商机制；对日志文件增加导出报告和可视化检索功能。总体而言，本项目已经完成课程要求中的国密算法综合应用，并在可视化、区块链存证、智能助手和篡改实验方面形成了较强的展示亮点。",
]


def iter_block_items(parent):
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def picture_count(paragraph):
    return len(paragraph._p.xpath(".//w:drawing")) + len(paragraph._p.xpath(".//w:pict"))


def set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def format_body(paragraph):
    paragraph.style = "Body Text" if "Body Text" in [s.name for s in paragraph.part.document.styles] else "Normal"
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def format_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)
        run.italic = True


def format_section_heading(paragraph):
    paragraph.style = "Heading 2"
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)


def next_description(paragraphs, start_index, current_heading_text):
    for p in paragraphs[start_index + 1 :]:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading") and text != current_heading_text:
            return ""
        if text.startswith("图 "):
            continue
        return text[:70]
    return ""


def add_existing_image_captions_and_number_headings(doc):
    paragraphs = list(doc.paragraphs)
    in_project = False
    current_num = None
    current_title = ""
    counts = {num: 0 for _, num in SECTION_ORDER}

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if text == "四 项目报告":
            in_project = True
            current_num = None
            continue
        if text == "五 项目总结":
            in_project = False
            current_num = None
            continue
        if not in_project:
            continue

        if p.style.name.startswith("Heading"):
            raw = text
            if raw in SECTION_NUM:
                current_num = SECTION_NUM[raw]
                current_title = raw
                set_paragraph_text(p, f"{current_num}. {raw}")
                format_section_heading(p)
            elif current_num is not None and raw:
                p.style = "Normal"
                format_body(p)
            continue

        pics = picture_count(p)
        if pics and current_num is not None:
            anchor = p
            desc = next_description(paragraphs, i, current_title)
            for _ in range(pics):
                counts[current_num] += 1
                caption_label = f"{current_title}演示" if current_title.endswith("功能") else f"{current_title}功能演示"
                caption_text = f"图 {current_num}-{counts[current_num]} {caption_label}"
                if desc:
                    caption_text += f"：{desc}"
                cap = insert_paragraph_after(anchor, caption_text, "Normal")
                format_caption(cap)
                anchor = cap


def remove_stray_project_heading_images(doc):
    for p in doc.paragraphs:
        if p.text.strip() == "四 项目报告":
            for node in list(p._p.xpath(".//w:drawing")) + list(p._p.xpath(".//w:pict")):
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)
            break


def remove_empty_heading_paragraphs(doc):
    for p in list(doc.paragraphs):
        if p.style.name.startswith("Heading") and not p.text.strip() and picture_count(p) == 0:
            parent = p._p.getparent()
            if parent is not None:
                parent.remove(p._p)


def add_section_explanations(doc):
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for num, title in SECTION_TITLE.items():
            if text == f"{num}. {title}":
                anchor = p
                for body in SECTION_EXPLANATIONS[num]:
                    anchor = insert_paragraph_after(anchor, body, "Body Text")
                    format_body(anchor)
                break


def add_architecture_block(doc):
    if not ARCH_IMAGE.exists():
        return
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "四 项目报告":
            anchor = p
            break
    if anchor is None:
        return

    h = insert_paragraph_after(anchor, "系统总体架构图", "Heading 2")
    format_section_heading(h)
    img_p = insert_paragraph_after(h, "", "Normal")
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(ARCH_IMAGE), width=Inches(6.4))
    cap = insert_paragraph_after(img_p, "图 0-1 系统总体架构：五层深度防御体系", "Normal")
    format_caption(cap)
    anchor = cap
    for body in ARCH_EXPLANATION:
        anchor = insert_paragraph_after(anchor, body, "Body Text")
        format_body(anchor)


def add_summary(doc):
    summary_heading = None
    source_heading = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "五 项目总结":
            summary_heading = p
        elif text == "六 程序源码":
            source_heading = p
            break
    if summary_heading is None:
        return

    # Remove stale empty paragraphs between summary and source only if they are blank.
    if source_heading is not None:
        for p in list(doc.paragraphs):
            if p._p is summary_heading._p:
                break

    anchor = summary_heading
    for body in SUMMARY_PARAGRAPHS:
        anchor = insert_paragraph_after(anchor, body, "Body Text")
        format_body(anchor)


def main():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))
    remove_stray_project_heading_images(doc)
    add_existing_image_captions_and_number_headings(doc)
    add_section_explanations(doc)
    add_architecture_block(doc)
    add_summary(doc)
    remove_empty_heading_paragraphs(doc)
    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()
